from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Plano de Criacao e Comercializacao - Sistema de Agentes IA Juridico.docx"


COLORS = {
    "heading": RGBColor(46, 116, 181),
    "heading_dark": RGBColor(31, 77, 120),
    "ink": RGBColor(11, 37, 69),
    "muted": RGBColor(89, 89, 89),
    "table_header": "F2F4F7",
    "callout": "F4F6F9",
    "risk": RGBColor(155, 28, 28),
}


REPLACEMENTS = {
    "Plano de Criacao e Comercializacao": "Plano de Criação e Comercialização",
    "Sistema de Agentes de IA para Escritorios de Advocacia": "Sistema de Agentes de IA para Escritórios de Advocacia",
    "juridico-operacional": "jurídico-operacional",
    "Prompt-Mestre": "Prompt-Mestre",
    "Criacao": "Criação",
    "criacao": "criação",
    "Comercializacao": "Comercialização",
    "comercializacao": "comercialização",
    "Juridico": "Jurídico",
    "juridico": "jurídico",
    "Juridica": "Jurídica",
    "juridica": "jurídica",
    "Decisao": "Decisão",
    "decisao": "decisão",
    "Estrategica": "Estratégica",
    "estrategica": "estratégica",
    "Vendavel": "Vendável",
    "vendavel": "vendável",
    "Escritorios": "Escritórios",
    "escritorios": "escritórios",
    "Escritorio": "Escritório",
    "escritorio": "escritório",
    "vendavel": "vendável",
    "implantavel": "implantável",
    "basicos": "básicos",
    "basico": "básico",
    "decisoes": "decisões",
    "validacao": "validação",
    "classificacao": "classificação",
    "Classificacao": "Classificação",
    "urgencia": "urgência",
    "informacoes": "informações",
    "proximos": "próximos",
    "recomendacao": "recomendação",
    "Recomendacao": "Recomendação",
    "gestao": "gestão",
    "visao": "visão",
    "organizacao": "organização",
    "padronizacao": "padronização",
    "automacao": "automação",
    "Automacao": "Automação",
    "automatizacao": "automatização",
    "substituicao": "substituição",
    "sao": "são",
    "peticoes": "petições",
    "Jurisprudencia": "Jurisprudência",
    "automatica": "automática",
    "automatico": "automático",
    "integracoes": "integrações",
    "Integracoes": "Integrações",
    "integracao": "integração",
    "Integracao": "Integração",
    "conciliacao": "conciliação",
    "confissoes": "confissões",
    "renuncias": "renúncias",
    "desistencias": "desistências",
    "avancado": "avançado",
    "bancaria": "bancária",
    "publicacao": "publicação",
    "permissoes": "permissões",
    "usuarios": "usuários",
    "Usuarios": "Usuários",
    "Conclusao": "Conclusão",
    "conclusao": "conclusão",
    "lancamento": "lançamento",
    "historico": "histórico",
    "negocio": "negócio",
    "area": "área",
    "notificacoes": "notificações",
    "geracao": "geração",
    "segregacao": "segregação",
    "experiencia": "experiência",
    "unica": "única",
    "auditavel": "auditável",
    "relatorios": "relatórios",
    "relatorio": "relatório",
    "producao": "produção",
    "comunicacao": "comunicação",
    "Politica": "Política",
    "politica": "política",
    "seguranca": "segurança",
    "pratica": "prática",
    "pratico": "prático",
    "tecnica": "técnica",
    "tecnico": "técnico",
    "criterios": "critérios",
    "Criterios": "Critérios",
    "pendencias": "pendências",
    "acoes": "ações",
    "minima": "mínima",
    "reducao": "redução",
    "Proximos": "Próximos",
    "modulos": "módulos",
    "historias": "histórias",
    "usuario": "usuário",
    "demonstracao": "demonstração",
    "versao": "versão",
    "Versao": "Versão",
    "alocacao": "alocação",
    "periodo": "período",
    "inadimplencia": "inadimplência",
    "responsaveis": "responsáveis",
    "responsavel": "responsável",
    "exportacao": "exportação",
    "expansao": "expansão",
    "aprovacao": "aprovação",
    "revisao": "revisão",
    "anonimização": "anonimização",
    "Nao ": "Não ",
    "nao ": "não ",
    " ate ": " até ",
    "Ha ": "Há ",
    " ha ": " há ",
    " esta ": " está ",
    ", e triado": ", é triado",
    "e lancar": "é lançar",
    "e validado": "é validado",
    "e nao": "e não",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["heading"], 16, 8),
        ("Heading 2", 13, COLORS["heading"], 12, 6),
        ("Heading 3", 12, COLORS["heading_dark"], 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Plano de Criacao e Comercializacao")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLORS["ink"]

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Sistema de Agentes de IA para Escritorios de Advocacia")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = COLORS["muted"]

    meta = doc.add_table(rows=4, cols=2)
    set_table_width(meta, [2300, 7060])
    rows = [
        ("Objetivo", "Transformar o Prompt-Mestre em um produto vendavel, seguro e implantavel."),
        ("Escopo recomendado", "MVP comercial com atendimento, triagem, CRM, documentos, minutas assistidas, prazos, compliance e indicadores basicos."),
        ("Estimativa para venda", "14 a 20 semanas para MVP pago com escopo controlado; 6 a 9 meses para SaaS robusto."),
        ("Premissa central", "A IA apoia, organiza e minuta; decisoes juridicas, envios externos e protocolos exigem validacao humana."),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], COLORS["table_header"])

    doc.add_paragraph()


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["callout"])
    clear_paragraph(cell.paragraphs[0])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = COLORS["ink"]
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def apply_text_replacements(doc):
    def replace_in_paragraph(paragraph):
        for run in paragraph.runs:
            text = run.text
            for source, target in REPLACEMENTS.items():
                text = text.replace(source, target)
            run.text = text

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = header
        set_cell_shading(cell, COLORS["table_header"])
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = COLORS["ink"]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    style_document(doc)
    add_title_block(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.add_run("Plano de criacao - Sistema de Agentes IA Juridico | Pagina ")
    add_page_number(footer)

    doc.add_heading("1. Decisao Estrategica", level=1)
    doc.add_paragraph(
        "O documento-base descreve uma arquitetura ampla de agentes para um escritorio de advocacia, "
        "incluindo atendimento, CRM, marketing juridico, peticionamento, controladoria, financeiro, "
        "qualidade, compliance, BI e gestao de conhecimento. Para transformar essa visao em produto, "
        "a recomendacao e lancar primeiro uma versao vendavel de menor risco, com foco em organizacao, "
        "triagem, rastreabilidade e minutas assistidas."
    )
    add_callout(
        doc,
        "Recomendacao principal",
        "Nao iniciar pelo sistema completo. Iniciar por um MVP juridico-operacional que resolva dores reais "
        "de escritorio: demora no atendimento, perda de leads, falta de padronizacao, documentos espalhados, "
        "pendencias sem controle e ausencia de rastreabilidade nas entregas da IA.",
    )

    doc.add_heading("2. Produto Vendavel Inicial", level=1)
    doc.add_paragraph(
        "O primeiro produto deve ser posicionado como uma plataforma de apoio operacional para escritorios "
        "de advocacia, com agentes de IA supervisionados. A promessa comercial deve ser produtividade, "
        "organizacao, padronizacao e seguranca institucional, sem prometer substituicao do advogado ou "
        "automatizacao integral de decisao juridica."
    )
    add_table(
        doc,
        ["Modulo", "Funcao no MVP", "Valor comercial"],
        [
            ("Atendimento e triagem", "Coletar dados, classificar urgencia e organizar o primeiro contato.", "Reduz tempo de resposta e evita perda de informacoes."),
            ("CRM juridico", "Gerenciar leads, etapas, follow-ups, propostas e status de contratacao.", "Aumenta conversao e previsibilidade comercial."),
            ("Clientes e processos", "Centralizar cadastro, historico, documentos, pendencias e responsaveis.", "Cria memoria operacional do escritorio."),
            ("Documentos", "Receber, classificar, nomear e vincular arquivos a cliente, processo ou tarefa.", "Reduz desorganizacao e retrabalho."),
            ("Minutas assistidas", "Gerar rascunhos de mensagens, checklists, propostas e pecas simples com validacao.", "Acelera producao sem dispensar revisao humana."),
            ("Prazos e tarefas", "Registrar pendencias, alertar vencimentos e distribuir responsaveis.", "Diminui risco operacional."),
            ("Compliance/LGPD", "Alertar dados sensiveis, risco etico, sigilo e necessidade de aprovacao.", "Aumenta confianca para uso profissional."),
            ("Indicadores basicos", "Mostrar leads, conversoes, tarefas, prazos, produtividade e inadimplencia simples.", "Ajuda o gestor a enxergar gargalos."),
        ],
        [2200, 4160, 3000],
    )

    doc.add_heading("3. Escopo que Deve Ficar Fora do MVP", level=1)
    doc.add_paragraph(
        "Alguns recursos sao importantes, mas aumentam muito o prazo, o risco tecnico ou o risco juridico. "
        "Eles devem entrar apenas depois que o produto estiver validado com clientes reais."
    )
    add_bullets(
        doc,
        [
            "Protocolo automatico de peticoes em tribunais.",
            "Integracao completa com PJe, e-SAJ, Projudi e demais sistemas processuais.",
            "Automacao de acordos, confissoes, renuncias ou desistencias.",
            "Jurisprudencia automatica sem fonte auditavel.",
            "Financeiro avancado com conciliacao bancaria completa.",
            "BI complexo com jurimetria profunda.",
            "Marketing juridico com publicacao automatica sem aprovacao humana.",
        ],
    )

    doc.add_heading("4. Fases de Desenvolvimento", level=1)
    add_table(
        doc,
        ["Fase", "Duracao", "Entregas principais", "Criterio de conclusao"],
        [
            ("0. Produto e requisitos", "1 a 2 semanas", "Escopo, personas, fluxos, regras de autonomia, arquitetura funcional e backlog.", "MVP fechado e priorizado."),
            ("1. Base tecnica", "2 a 3 semanas", "Login, perfis, banco de dados, auditoria, cadastros centrais e estrutura dos agentes.", "Ambiente navegavel com usuarios e dados essenciais."),
            ("2. Nucleo operacional", "4 a 6 semanas", "Clientes, leads, CRM, atendimentos, documentos, tarefas, prazos e historico.", "Escritorio consegue operar uma rotina basica no sistema."),
            ("3. Agentes do MVP", "6 a 8 semanas", "Coordenador, triagem, atendimento, CRM, documentos, minutas, revisao, prazos e compliance.", "Agentes geram entregas padronizadas com fontes, riscos e validacao."),
            ("4. Integracoes essenciais", "3 a 6 semanas", "E-mail, agenda, armazenamento, exportacao de relatorios e canais de atendimento.", "Fluxos principais conectados a ferramentas reais."),
            ("5. Seguranca e governanca", "3 a 4 semanas", "LGPD, logs, permissoes, termos, backups, bloqueios e matriz de risco.", "Produto apto a tratar dados sensiveis com controle."),
            ("6. Piloto", "4 a 6 semanas", "Uso em 2 a 5 escritorios, ajustes, metricas, treinamento e suporte.", "Produto validado por usuarios reais."),
            ("7. Comercializacao", "2 a 4 semanas", "Site, demo, contratos, planos, onboarding, materiais comerciais e suporte.", "Produto pronto para venda recorrente."),
        ],
        [1500, 1300, 4260, 2300],
    )

    doc.add_heading("5. Cronograma Recomendado", level=1)
    add_table(
        doc,
        ["Marco", "Quando", "Resultado esperado"],
        [
            ("Demonstração navegavel", "Semana 6 a 8", "Sistema com telas principais, fluxo de atendimento/CRM e agentes em modo demonstracao."),
            ("MVP interno", "Semana 10 a 12", "Equipe consegue cadastrar clientes, rodar triagens, gerar minutas e controlar pendencias."),
            ("MVP pago controlado", "Semana 14 a 20", "Produto apto a vender para primeiros clientes com onboarding acompanhado."),
            ("SaaS robusto", "Mes 6 a 9", "Produto mais estavel, com integracoes, relatorios, seguranca reforcada e base de clientes inicial."),
            ("Produto completo ampliado", "Mes 9 a 12+", "Expansao para modulos avancados: financeiro, BI, controladoria profunda e integracoes processuais."),
        ],
        [2500, 1800, 5060],
    )

    doc.add_heading("6. Arquitetura Funcional", level=1)
    doc.add_paragraph(
        "A arquitetura deve separar a plataforma operacional dos agentes de IA. O sistema registra dados, "
        "controla permissoes, organiza historico e aplica regras de negocio. Os agentes atuam sobre esse "
        "contexto, sempre com limites de autonomia e rastreabilidade."
    )
    add_numbered(
        doc,
        [
            "Camada de interface: dashboard, CRM, clientes, processos, documentos, tarefas, relatorios e area dos agentes.",
            "Camada operacional: regras de negocio, permissoes, auditoria, notificacoes, filas de revisao e workflow.",
            "Camada de dados: clientes, leads, processos, documentos, atendimentos, propostas, tarefas, prazos, pagamentos e logs.",
            "Camada de IA: prompts versionados, agentes especializados, base de conhecimento, classificadores e geracao de minutas.",
            "Camada de seguranca: LGPD, segregacao por escritorio, criptografia, backups, logs, controle de acesso e alertas de risco.",
            "Camada de integracoes: e-mail, agenda, armazenamento, atendimento e futuras integracoes processuais.",
        ],
    )

    doc.add_heading("7. Agentes da Primeira Versao", level=1)
    add_table(
        doc,
        ["Agente", "Prioridade", "Autonomia", "Saida principal"],
        [
            ("Coordenador Geral", "Alta", "Classifica e roteia; nao decide juridicamente.", "Relatorio de demanda com riscos e proximos passos."),
            ("Triagem e Classificacao", "Alta", "Organiza dados e identifica urgencia.", "Tipo de demanda, area, documentos e setor indicado."),
            ("Atendimento Inicial", "Alta", "Coleta informacoes e sugere mensagens.", "Mensagem humanizada e formulario de triagem."),
            ("CRM e Relacionamento", "Alta", "Sugere follow-ups e status.", "Registro de lead, etapa, pendencias e risco de perda."),
            ("Gestao Documental", "Alta", "Classifica e nomeia arquivos.", "Vinculo, nome padronizado e pendencias documentais."),
            ("Prazos e Agenda", "Alta", "Alerta e organiza tarefas.", "Lista de prazos, responsaveis e riscos."),
            ("Peticionamento Assistido", "Media", "Gera rascunho, nunca versao final.", "Minuta com lacunas, fontes e alerta de revisao."),
            ("Revisao Juridica", "Media", "Aponta inconsistencias.", "Checklist de revisao tecnica e riscos."),
            ("Compliance, Etica e LGPD", "Alta", "Pode bloquear saida sensivel.", "Alerta de risco, gravidade e providencia recomendada."),
            ("BI Basico", "Media", "Analisa dados cadastrados.", "Indicadores simples e gargalos operacionais."),
        ],
        [2600, 1300, 2300, 3160],
    )

    doc.add_heading("8. Equipe Recomendada", level=1)
    add_table(
        doc,
        ["Perfil", "Alocacao sugerida", "Responsabilidade"],
        [
            ("Product owner juridico", "Meio periodo", "Validar fluxos, linguagem, riscos, regras de autonomia e requisitos do escritorio."),
            ("Tech lead/full-stack senior", "Integral", "Arquitetura, backend, frontend, seguranca tecnica e revisao de codigo."),
            ("Desenvolvedor full-stack", "Integral", "Telas, APIs, cadastros, workflows e integracoes."),
            ("Especialista IA/prompt engineering", "Meio a integral", "Agentes, prompts, avaliacoes, guardrails e rastreabilidade."),
            ("UX/UI designer", "Parcial", "Fluxos, telas, experiencia do advogado, dashboard e onboarding."),
            ("QA/testes", "Parcial", "Testes funcionais, regressao, seguranca basica e aceitacao."),
            ("Consultor LGPD/compliance", "Pontual", "Politicas, termos, riscos de dados e governanca."),
        ],
        [2600, 2100, 4660],
    )

    doc.add_heading("9. Criterios para o Sistema Estar Apto a Ser Vendido", level=1)
    add_bullets(
        doc,
        [
            "O sistema resolve um fluxo completo: lead entra, e triado, vira oportunidade, recebe follow-up, gera minuta e cria pendencias.",
            "Todas as respostas relevantes da IA mostram informacoes usadas, lacunas, riscos e necessidade de validacao humana.",
            "Usuarios conseguem operar sem treinamento longo.",
            "Ha controle de acesso por perfil e segregacao de dados por escritorio.",
            "Ha logs de auditoria das acoes e das respostas dos agentes.",
            "O sistema possui termos, politica de privacidade, regras de uso e avisos de responsabilidade profissional.",
            "Ha backup, rotina minima de suporte e procedimento de incidentes.",
            "O piloto demonstrou ganho pratico em tempo de resposta, organizacao e reducao de retrabalho.",
        ],
    )

    doc.add_heading("10. Riscos e Controles", level=1)
    add_table(
        doc,
        ["Risco", "Impacto", "Controle recomendado"],
        [
            ("IA inventar fatos, jurisprudencia ou documentos", "Alto", "Rastreabilidade obrigatoria, respostas com lacunas e bloqueio de afirmacoes sem fonte."),
            ("Uso indevido em comunicacao juridica sensivel", "Alto", "Fila de aprovacao humana para mensagens, propostas, peticoes, contratos e pareceres."),
            ("Exposicao de dados pessoais ou sigilosos", "Alto", "Permissoes, segregacao por escritorio, criptografia, anonimização e logs."),
            ("Escopo grande demais para primeira versao", "Alto", "MVP limitado, backlog faseado e criterio claro de nao incluir integracoes processuais no inicio."),
            ("Baixa confianca dos advogados", "Medio/Alto", "Explicar fontes, permitir edicao, registrar historico e usar pilotos com feedback real."),
            ("Produto dificil de vender", "Medio", "Vender dor objetiva: atendimento, CRM, documentos, tarefas e minutas assistidas."),
        ],
        [2800, 1600, 4960],
    )

    doc.add_heading("11. Pacotes Comerciais Sugeridos", level=1)
    add_table(
        doc,
        ["Plano", "Publico", "Recursos"],
        [
            ("Starter", "Escritorios pequenos", "Atendimento, triagem, CRM, documentos, tarefas e agentes basicos."),
            ("Professional", "Escritorios em crescimento", "Tudo do Starter, minutas assistidas, compliance, indicadores e integracoes essenciais."),
            ("Enterprise", "Operacoes maiores", "Permissoes avancadas, auditoria ampliada, customizacoes, suporte prioritario e BI expandido."),
        ],
        [2000, 2600, 4760],
    )

    doc.add_heading("12. Proximos Passos Imediatos", level=1)
    add_numbered(
        doc,
        [
            "Fechar o escopo do MVP em ate 10 modulos e ate 10 agentes iniciais.",
            "Definir stack tecnica, modelo de hospedagem e politica de dados.",
            "Desenhar as telas principais: dashboard, CRM, cliente, processo, documentos, tarefas e agentes.",
            "Criar backlog com historias de usuario e criterios de aceite.",
            "Construir demonstracao em 6 a 8 semanas.",
            "Selecionar 2 a 5 escritorios para piloto controlado.",
            "Preparar materiais comerciais e contrato SaaS antes da venda paga.",
        ],
    )

    doc.add_heading("Conclusao", level=1)
    doc.add_paragraph(
        "O sistema pode se tornar um produto vendavel em aproximadamente 14 a 20 semanas, desde que "
        "o primeiro lancamento seja tratado como MVP juridico-operacional e nao como automacao completa "
        "de escritorio. A vantagem competitiva inicial esta em unir atendimento, CRM, documentos, tarefas, "
        "minutas assistidas e compliance em uma unica experiencia simples, auditavel e segura."
    )

    apply_text_replacements(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build()
